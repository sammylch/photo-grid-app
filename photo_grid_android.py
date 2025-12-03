"""
Photo Grid Generator - Android Version
Kivy-based mobile app for creating 2x2 photo grids with captions
"""

from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.button import Button
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.scrollview import ScrollView
from kivy.uix.image import Image as KivyImage
from kivy.uix.popup import Popup
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.spinner import Spinner
from kivy.core.window import Window
from kivy.graphics import Color, Rectangle
from kivy.utils import platform

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import io
import json
from datetime import datetime

# Android permissions
if platform == 'android':
    from android.permissions import request_permissions, Permission
    request_permissions([
        Permission.READ_EXTERNAL_STORAGE,
        Permission.WRITE_EXTERNAL_STORAGE,
        Permission.CAMERA
    ])


class PhotoGridCell(BoxLayout):
    """Individual photo grid cell with image and caption"""
    def __init__(self, index, callback, remove_callback, **kwargs):
        super().__init__(**kwargs)
        self.orientation = 'vertical'
        self.index = index
        self.callback = callback
        self.remove_callback = remove_callback
        self.photo_path = None
        
        # Image display
        self.image = KivyImage(
            source='',
            allow_stretch=True,
            keep_ratio=True,
            size_hint=(1, 0.65)
        )
        
        # Add background
        with self.image.canvas.before:
            Color(0.9, 0.9, 0.9, 1)
            self.rect = Rectangle(size=self.image.size, pos=self.image.pos)
        
        self.image.bind(size=self._update_rect, pos=self._update_rect)
        
        # Select button
        select_btn = Button(
            text='Select Photo',
            size_hint=(1, 0.12),
            on_press=lambda x: self.callback(self.index)
        )
        
        # Caption and delete button layout
        caption_layout = BoxLayout(size_hint=(1, 0.12), spacing=5)
        
        self.caption_input = TextInput(
            text=f'Photo {index + 1}',
            multiline=False,
            size_hint=(0.75, 1)
        )
        
        delete_btn = Button(
            text='✕ Del',
            size_hint=(0.25, 1),
            background_color=(1, 0.3, 0.3, 1),
            on_press=lambda x: self.remove_callback(self.index)
        )
        
        caption_layout.add_widget(self.caption_input)
        caption_layout.add_widget(delete_btn)
        
        self.add_widget(self.image)
        self.add_widget(select_btn)
        self.add_widget(caption_layout)
    
    def _update_rect(self, instance, value):
        self.rect.pos = instance.pos
        self.rect.size = instance.size
    
    def set_photo(self, path):
        """Set photo for this cell"""
        self.photo_path = path
        self.image.source = path
        self.image.reload()
    
    def get_caption(self):
        """Get caption text"""
        return self.caption_input.text
    
    def set_caption(self, text):
        """Set caption text"""
        self.caption_input.text = text


class PhotoGridApp(App):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.photos = []
        self.current_page = 0
        self.page_photos = {}
        self.page_titles = {}
        self.page_captions = {}
        self.photos_per_page = 4
        self.grid_layout_type = '2x2'  # Default to 2x2
        
        # Settings
        self.title_font_size = 14
        self.title_bold = True
        self.title_underline = True
        self.caption_font_size = 10
        self.caption_bold = False
        self.caption_underline = False
        self.caption_prefix = 'Photo'
        self.header_text = ''
        self.footer_line1 = 'File Ref. :'
        self.footer_line2 = f'Photo taken on {datetime.now().strftime("%d.%m.%Y")}'
        
        # Get storage path
        if platform == 'android':
            from android.storage import primary_external_storage_path
            self.storage_path = primary_external_storage_path()
        else:
            self.storage_path = os.path.expanduser('~')
        
        self.settings_file = os.path.join(self.storage_path, 'photo_grid_settings.json')
        self.last_photo_path = self.storage_path  # Remember last folder
        self.initializing = True  # Flag to prevent handlers during init
        self.load_settings()
    
    def build(self):
        """Build the main UI with modern design"""
        self.title = 'Photo Grid Generator'
        
        # Set window background color
        Window.clearcolor = (0.95, 0.95, 0.97, 1)
        
        # Main layout with modern styling
        main_layout = BoxLayout(orientation='vertical', padding=5, spacing=8)
        
        # Top toolbar with modern gradient-like colors
        toolbar = BoxLayout(size_hint=(1, 0.08), spacing=4)
        
        add_btn = Button(
            text='📷 Add',
            on_press=self.show_file_chooser,
            background_color=(0.26, 0.59, 0.98, 1),  # Material Blue
            background_normal='',
            color=(1, 1, 1, 1),
            bold=True
        )
        save_word_btn = Button(
            text='� WSave',
            on_press=self.save_word,
            background_color=(0.30, 0.69, 0.31, 1),  # Material Green
            background_normal='',
            color=(1, 1, 1, 1),
            bold=True
        )
        reset_btn = Button(
            text='� PReset',
            on_press=self.reset_project,
            background_color=(0.96, 0.26, 0.21, 1),  # Material Red
            background_normal='',
            color=(1, 1, 1, 1),
            bold=True
        )
        
        toolbar.add_widget(add_btn)
        toolbar.add_widget(save_word_btn)
        toolbar.add_widget(reset_btn)
        
        main_layout.add_widget(toolbar)
        
        # Page navigation with modern styling
        nav_layout = BoxLayout(size_hint=(1, 0.06), spacing=4)
        
        prev_btn = Button(
            text='◀ Prev',
            on_press=lambda x: self.change_page(-1),
            background_color=(0.38, 0.49, 0.55, 1),  # Material Blue Grey
            background_normal='',
            color=(1, 1, 1, 1)
        )
        self.page_label = Label(
            text='Page 1 of 1',
            color=(0.2, 0.2, 0.2, 1),
            bold=True
        )
        next_btn = Button(
            text='Next ▶',
            on_press=lambda x: self.change_page(1),
            background_color=(0.38, 0.49, 0.55, 1),  # Material Blue Grey
            background_normal='',
            color=(1, 1, 1, 1)
        )
        
        nav_layout.add_widget(prev_btn)
        nav_layout.add_widget(self.page_label)
        nav_layout.add_widget(next_btn)
        
        main_layout.add_widget(nav_layout)
        
        # Combined title and caption prefix row with modern styling
        title_caption_layout = BoxLayout(size_hint=(1, 0.06), spacing=4)
        
        # Title section (left side)
        title_caption_layout.add_widget(Label(
            text='Title:',
            size_hint=(0.12, 1),
            color=(0.2, 0.2, 0.2, 1),
            bold=True
        ))
        self.title_input = TextInput(
            text='Title',
            multiline=False,
            size_hint=(0.38, 1),
            background_color=(1, 1, 1, 1),
            foreground_color=(0.2, 0.2, 0.2, 1),
            cursor_color=(0.26, 0.59, 0.98, 1),
            padding=[8, 8]
        )
        self.title_input.bind(text=self.on_title_change)
        title_caption_layout.add_widget(self.title_input)
        
        # Caption prefix section (right side)
        title_caption_layout.add_widget(Label(
            text='Prefix:',
            size_hint=(0.12, 1),
            color=(0.2, 0.2, 0.2, 1),
            bold=True
        ))
        self.caption_prefix_input = TextInput(
            text=self.caption_prefix,
            multiline=False,
            size_hint=(0.38, 1),
            background_color=(1, 1, 1, 1),
            foreground_color=(0.2, 0.2, 0.2, 1),
            cursor_color=(0.26, 0.59, 0.98, 1),
            padding=[8, 8]
        )
        self.caption_prefix_input.bind(text=self.on_caption_prefix_change)
        title_caption_layout.add_widget(self.caption_prefix_input)
        
        main_layout.add_widget(title_caption_layout)
        
        # Grid layout selector with modern styling
        grid_selector_layout = BoxLayout(size_hint=(1, 0.06), spacing=4)
        grid_selector_layout.add_widget(Label(
            text='Layout:',
            size_hint=(0.25, 1),
            color=(0.2, 0.2, 0.2, 1),
            bold=True
        ))
        self.grid_spinner = Spinner(
            text='2x2',
            values=('2x1', '2x2'),
            size_hint=(0.75, 1),
            background_color=(0.26, 0.59, 0.98, 1),
            background_normal='',
            color=(1, 1, 1, 1)
        )
        self.grid_spinner.bind(text=self.change_layout)
        grid_selector_layout.add_widget(self.grid_spinner)
        
        main_layout.add_widget(grid_selector_layout)
        
        # Photo grid (2x2)
        self.grid_container = GridLayout(cols=2, spacing=10, size_hint=(1, 0.7))
        self.grid_cells = []
        
        for i in range(4):
            cell = PhotoGridCell(i, self.select_photo_for_cell, self.remove_photo_from_cell)
            self.grid_cells.append(cell)
            self.grid_container.add_widget(cell)
        
        main_layout.add_widget(self.grid_container)
        
        # Settings button with modern styling
        settings_btn = Button(
            text='⚙ Settings',
            size_hint=(1, 0.06),
            on_press=self.show_settings,
            background_color=(0.61, 0.35, 0.71, 1),  # Material Purple
            background_normal='',
            color=(1, 1, 1, 1),
            bold=True
        )
        main_layout.add_widget(settings_btn)
        
        # Sync spinner with loaded settings
        self.grid_spinner.text = self.grid_layout_type
        
        # Initialization complete - enable handlers
        self.initializing = False
        
        return main_layout
    
    def change_layout(self, spinner, text):
        """Change grid layout between 2x1 and 2x2 and reorganize photos"""
        old_photos_per_page = self.photos_per_page
        
        # Update grid_layout_type to match spinner
        self.grid_layout_type = text
        print(f"DEBUG: Layout changed to: {text}")
        
        # Clear existing grid
        self.grid_container.clear_widgets()
        self.grid_cells = []
        
        if text == '2x1':
            self.grid_container.cols = 1
            self.photos_per_page = 2
            num_cells = 2
        else:
            self.grid_container.cols = 2
            self.photos_per_page = 4
            num_cells = 4
        
        # Reorganize photos if layout changed
        if old_photos_per_page != self.photos_per_page:
            # Collect all photos in order
            all_photos_ordered = []
            old_page_photos = dict(self.page_photos)
            old_page_captions = dict(self.page_captions)
            
            # Sort by page and position to maintain order
            sorted_keys = sorted(old_page_photos.keys(), 
                               key=lambda x: (int(x.split('_')[0]), int(x.split('_')[1])))
            
            for key in sorted_keys:
                photo_path = old_page_photos[key]
                caption = old_page_captions.get(key, '')
                all_photos_ordered.append((photo_path, caption))
            
            # Clear old mappings
            self.page_photos.clear()
            self.page_captions.clear()
            
            # Redistribute photos with new layout
            for idx, (photo_path, caption) in enumerate(all_photos_ordered):
                page_num = idx // self.photos_per_page
                grid_idx = idx % self.photos_per_page
                page_key = f'{page_num}_{grid_idx}'
                
                self.page_photos[page_key] = photo_path
                if caption:
                    self.page_captions[page_key] = caption
            
            # Reset to first page
            self.current_page = 0
        
        # Rebuild grid
        for i in range(num_cells):
            cell = PhotoGridCell(i, self.select_photo_for_cell, self.remove_photo_from_cell)
            self.grid_cells.append(cell)
            self.grid_container.add_widget(cell)
        
        self.update_preview()
    
    def on_title_change(self, instance, value):
        """Save title when it changes"""
        self.page_titles[self.current_page] = value
    
    def on_caption_prefix_change(self, instance, value):
        """Update caption prefix and all captions when it changes"""
        # Don't run during initialization
        if hasattr(self, 'initializing') and self.initializing:
            return
        
        new_prefix = value.strip()
        if new_prefix and new_prefix != self.caption_prefix and len(new_prefix) > 0:
            self.caption_prefix = new_prefix
            try:
                if self.page_photos:
                    self.update_all_captions_with_prefix()
                    self.update_preview()
                self.save_settings()
            except Exception as e:
                print(f'Error updating caption prefix: {e}')
    
    def change_page(self, direction):
        """Navigate between pages"""
        max_page = max(0, (len(self.photos) + self.photos_per_page - 1) // self.photos_per_page - 1)
        
        new_page = self.current_page + direction
        if 0 <= new_page <= max_page:
            self.current_page = new_page
            self.update_preview()
    
    def update_preview(self):
        """Update the grid preview for current page"""
        max_page = max(0, (len(self.photos) + self.photos_per_page - 1) // self.photos_per_page - 1)
        self.page_label.text = f'Page {self.current_page + 1} of {max_page + 1}'
        
        # Update title
        if self.current_page in self.page_titles:
            self.title_input.text = self.page_titles[self.current_page]
        else:
            self.title_input.text = 'Title'
        
        # Update grid cells
        for i in range(len(self.grid_cells)):
            page_key = f'{self.current_page}_{i}'
            
            if page_key in self.page_photos:
                # Photo exists in this position
                self.grid_cells[i].set_photo(self.page_photos[page_key])
                
                # Set caption - use saved caption or generate default
                if page_key in self.page_captions:
                    self.grid_cells[i].set_caption(self.page_captions[page_key])
                else:
                    # Generate default caption
                    start_idx = self.current_page * self.photos_per_page
                    default_caption = f'{self.caption_prefix} {start_idx + i + 1}'
                    self.page_captions[page_key] = default_caption
                    self.grid_cells[i].set_caption(default_caption)
            else:
                # No photo assigned to this position yet
                start_idx = self.current_page * self.photos_per_page
                if start_idx + i < len(self.photos):
                    photo_path = self.photos[start_idx + i]
                    self.page_photos[page_key] = photo_path
                    self.grid_cells[i].set_photo(photo_path)
                    
                    # Generate caption with correct sequential number
                    photo_num = start_idx + i + 1
                    default_caption = f'{self.caption_prefix} {photo_num}'
                    self.page_captions[page_key] = default_caption
                    self.grid_cells[i].set_caption(default_caption)
                else:
                    # Empty cell
                    self.grid_cells[i].set_photo('')
                    self.grid_cells[i].set_caption('')
    
    def select_photo_for_cell(self, cell_index):
        """Show file chooser for selecting photo"""
        self.selected_cell_index = cell_index
        self.show_file_chooser()
    
    def update_all_captions_with_prefix(self):
        """Update all captions to use new prefix while maintaining numbers"""
        if not self.page_photos:
            return
        
        # Sort all page_photos by page and position to get correct order
        try:
            sorted_keys = sorted(self.page_photos.keys(), 
                               key=lambda x: (int(x.split('_')[0]), int(x.split('_')[1])))
            
            # Update each caption with new prefix, keeping the sequential number
            for idx, page_key in enumerate(sorted_keys):
                photo_num = idx + 1
                new_caption = f'{self.caption_prefix} {photo_num}'
                self.page_captions[page_key] = new_caption
        except Exception as e:
            print(f'Error updating captions: {e}')
            # Fallback: just update visible captions
            for i in range(len(self.grid_cells)):
                page_key = f'{self.current_page}_{i}'
                if page_key in self.page_photos:
                    self.page_captions[page_key] = f'{self.caption_prefix} {i + 1}'
    
    def remove_photo_from_cell(self, cell_index):
        """Remove photo from specific cell permanently"""
        page_key = f'{self.current_page}_{cell_index}'
        if page_key in self.page_photos:
            # Show confirmation
            content = BoxLayout(orientation='vertical', padding=10, spacing=10)
            content.add_widget(Label(text='Remove this photo from the grid?\n(Will not reappear on page navigation)'))
            
            button_layout = BoxLayout(size_hint=(1, 0.3), spacing=10)
            yes_btn = Button(text='Yes', background_color=(1, 0.3, 0.3, 1))
            no_btn = Button(text='No')
            button_layout.add_widget(yes_btn)
            button_layout.add_widget(no_btn)
            content.add_widget(button_layout)
            
            popup = Popup(title='Remove Photo', content=content, size_hint=(0.8, 0.35))
            
            def on_yes(instance):
                # Remove from page_photos and page_captions
                photo_path = self.page_photos[page_key]
                del self.page_photos[page_key]
                if page_key in self.page_captions:
                    del self.page_captions[page_key]
                
                # Also remove from global photos list to prevent reappearing
                if photo_path in self.photos:
                    self.photos.remove(photo_path)
                
                # Clear the cell
                self.grid_cells[cell_index].set_photo('')
                self.grid_cells[cell_index].set_caption('')
                popup.dismiss()
            
            def on_no(instance):
                popup.dismiss()
            
            yes_btn.bind(on_press=on_yes)
            no_btn.bind(on_press=on_no)
            
            popup.open()
    
    def show_file_chooser(self, *args):
        """Show file chooser dialog with multiple selection"""
        content = BoxLayout(orientation='vertical')
        
        file_chooser = FileChooserIconView(
            path=self.last_photo_path,  # Remember last folder
            filters=['*.jpg', '*.jpeg', '*.png', '*.bmp', '*.gif'],
            multiselect=True  # Enable multiple selection
        )
        
        button_layout = BoxLayout(size_hint=(1, 0.1), spacing=10)
        select_btn = Button(text='Select')
        cancel_btn = Button(text='Cancel')
        
        button_layout.add_widget(select_btn)
        button_layout.add_widget(cancel_btn)
        
        content.add_widget(file_chooser)
        content.add_widget(button_layout)
        
        popup = Popup(
            title='Select Photos (Multiple)',
            content=content,
            size_hint=(0.9, 0.9)
        )
        
        def on_select(instance):
            if file_chooser.selection:
                # Remember the folder for next time
                if file_chooser.selection:
                    self.last_photo_path = os.path.dirname(file_chooser.selection[0])
                
                if hasattr(self, 'selected_cell_index'):
                    # Single cell selection
                    photo_path = file_chooser.selection[0]
                    page_key = f'{self.current_page}_{self.selected_cell_index}'
                    self.page_photos[page_key] = photo_path
                    if photo_path not in self.photos:
                        self.photos.append(photo_path)
                    self.grid_cells[self.selected_cell_index].set_photo(photo_path)
                    delattr(self, 'selected_cell_index')  # Clear selection
                else:
                    # Bulk add - multiple photos
                    for photo_path in file_chooser.selection:
                        if photo_path not in self.photos:
                            self.photos.append(photo_path)
                    self.update_preview()
            popup.dismiss()
        
        def on_cancel(instance):
            if hasattr(self, 'selected_cell_index'):
                delattr(self, 'selected_cell_index')  # Clear selection
            popup.dismiss()
        
        select_btn.bind(on_press=on_select)
        cancel_btn.bind(on_press=on_cancel)
        
        popup.open()
    
    def show_settings(self, *args):
        """Show settings dialog with font options"""
        from kivy.uix.checkbox import CheckBox
        
        content = BoxLayout(orientation='vertical', padding=10, spacing=5)
        scroll = ScrollView(size_hint=(1, 0.85))
        scroll_content = BoxLayout(orientation='vertical', spacing=5, size_hint_y=None)
        scroll_content.bind(minimum_height=scroll_content.setter('height'))
        
        # Info label
        info_label = Label(
            text='[b]Settings (Auto-saved)[/b]',
            markup=True,
            size_hint=(1, None),
            height=30
        )
        scroll_content.add_widget(info_label)
        
        # Title Font Settings
        title_font_label = Label(
            text='[b]Title Font[/b]',
            markup=True,
            size_hint=(1, None),
            height=30
        )
        scroll_content.add_widget(title_font_label)
        
        # Title font size
        title_size_layout = BoxLayout(size_hint=(1, None), height=35)
        title_size_layout.add_widget(Label(text='Size:', size_hint=(0.3, 1)))
        title_size_spinner = Spinner(
            text=str(self.title_font_size),
            values=[str(i) for i in range(10, 25)],
            size_hint=(0.7, 1)
        )
        title_size_layout.add_widget(title_size_spinner)
        scroll_content.add_widget(title_size_layout)
        
        # Title bold/underline
        title_style_layout = BoxLayout(size_hint=(1, None), height=35)
        title_bold_check = CheckBox(active=self.title_bold, size_hint=(0.15, 1))
        title_style_layout.add_widget(title_bold_check)
        title_style_layout.add_widget(Label(text='Bold', size_hint=(0.35, 1)))
        title_underline_check = CheckBox(active=self.title_underline, size_hint=(0.15, 1))
        title_style_layout.add_widget(title_underline_check)
        title_style_layout.add_widget(Label(text='Underline', size_hint=(0.35, 1)))
        scroll_content.add_widget(title_style_layout)
        
        # Caption Font Settings
        caption_font_label = Label(
            text='[b]Caption Font[/b]',
            markup=True,
            size_hint=(1, None),
            height=30
        )
        scroll_content.add_widget(caption_font_label)
        
        # Caption font size
        caption_size_layout = BoxLayout(size_hint=(1, None), height=35)
        caption_size_layout.add_widget(Label(text='Size:', size_hint=(0.3, 1)))
        caption_size_spinner = Spinner(
            text=str(self.caption_font_size),
            values=[str(i) for i in range(6, 16)],
            size_hint=(0.7, 1)
        )
        caption_size_layout.add_widget(caption_size_spinner)
        scroll_content.add_widget(caption_size_layout)
        
        # Caption bold/underline
        caption_style_layout = BoxLayout(size_hint=(1, None), height=35)
        caption_bold_check = CheckBox(active=self.caption_bold, size_hint=(0.15, 1))
        caption_style_layout.add_widget(caption_bold_check)
        caption_style_layout.add_widget(Label(text='Bold', size_hint=(0.35, 1)))
        caption_underline_check = CheckBox(active=self.caption_underline, size_hint=(0.15, 1))
        caption_style_layout.add_widget(caption_underline_check)
        caption_style_layout.add_widget(Label(text='Underline', size_hint=(0.35, 1)))
        scroll_content.add_widget(caption_style_layout)
        
        # Separator
        scroll_content.add_widget(Label(text='─' * 40, size_hint=(1, None), height=20))
        
        # Note about caption prefix
        prefix_note = Label(
            text='[i]Caption Prefix can be edited in the main screen\nnext to the Title field[/i]',
            markup=True,
            size_hint=(1, None),
            height=40,
            color=(0.7, 0.7, 0.7, 1)
        )
        scroll_content.add_widget(prefix_note)
        
        # Header
        header_layout = BoxLayout(size_hint=(1, None), height=35)
        header_layout.add_widget(Label(text='Header:', size_hint=(0.4, 1)))
        header_input = TextInput(text=self.header_text, multiline=False, size_hint=(0.6, 1))
        header_layout.add_widget(header_input)
        scroll_content.add_widget(header_layout)
        
        # Footer line 1
        footer1_layout = BoxLayout(size_hint=(1, None), height=35)
        footer1_layout.add_widget(Label(text='Footer Line 1:', size_hint=(0.4, 1)))
        footer1_input = TextInput(text=self.footer_line1, multiline=False, size_hint=(0.6, 1))
        footer1_layout.add_widget(footer1_input)
        scroll_content.add_widget(footer1_layout)
        
        # Footer line 2
        footer2_layout = BoxLayout(size_hint=(1, None), height=35)
        footer2_layout.add_widget(Label(text='Footer Line 2:', size_hint=(0.4, 1)))
        footer2_input = TextInput(text=self.footer_line2, multiline=False, size_hint=(0.6, 1))
        footer2_layout.add_widget(footer2_input)
        scroll_content.add_widget(footer2_layout)
        
        scroll.add_widget(scroll_content)
        content.add_widget(scroll)
        
        # Buttons
        button_layout = BoxLayout(size_hint=(1, 0.15), spacing=10)
        apply_btn = Button(text='✓ Apply & Close', background_color=(0.2, 0.8, 0.2, 1))
        cancel_btn = Button(text='✕ Cancel')
        button_layout.add_widget(apply_btn)
        button_layout.add_widget(cancel_btn)
        content.add_widget(button_layout)
        
        popup = Popup(title='⚙ Settings', content=content, size_hint=(0.95, 0.85))
        
        def on_apply(instance):
            # Check if caption prefix changed (from main screen input)
            old_prefix = self.caption_prefix
            new_prefix = self.caption_prefix_input.text.strip()
            prefix_changed = (old_prefix != new_prefix)
            
            # Update settings
            self.title_font_size = int(title_size_spinner.text)
            self.title_bold = title_bold_check.active
            self.title_underline = title_underline_check.active
            self.caption_font_size = int(caption_size_spinner.text)
            self.caption_bold = caption_bold_check.active
            self.caption_underline = caption_underline_check.active
            self.caption_prefix = new_prefix
            self.header_text = header_input.text
            self.footer_line1 = footer1_input.text
            self.footer_line2 = footer2_input.text
            
            # Update all captions if prefix changed
            if prefix_changed and self.page_photos:
                self.update_all_captions_with_prefix()
            
            self.save_settings()
            self.update_preview()  # Refresh display
            self.show_message('Settings Saved', 'Your settings have been saved!')
            popup.dismiss()
        
        def on_cancel(instance):
            popup.dismiss()
        
        apply_btn.bind(on_press=on_apply)
        cancel_btn.bind(on_press=on_cancel)
        
        popup.open()
    
    def reset_project(self, *args):
        """Reset all photos and settings"""
        content = BoxLayout(orientation='vertical', padding=10, spacing=10)
        content.add_widget(Label(text='Reset all photos and start fresh?'))
        
        button_layout = BoxLayout(size_hint=(1, 0.3), spacing=10)
        yes_btn = Button(text='Yes')
        no_btn = Button(text='No')
        button_layout.add_widget(yes_btn)
        button_layout.add_widget(no_btn)
        content.add_widget(button_layout)
        
        popup = Popup(title='Reset Project', content=content, size_hint=(0.8, 0.3))
        
        def on_yes(instance):
            self.photos = []
            self.page_photos = {}
            self.page_titles = {}
            self.page_captions = {}
            self.current_page = 0
            self.update_preview()
            popup.dismiss()
        
        def on_no(instance):
            popup.dismiss()
        
        yes_btn.bind(on_press=on_yes)
        no_btn.bind(on_press=on_no)
        
        popup.open()
    
    def save_word(self, *args):
        """Generate and save Word document with custom filename"""
        if not self.photos:
            self.show_message('No Photos', 'Please add photos first!')
            return
        
        # Show filename input dialog
        self.show_save_dialog('word')
    

    def show_save_dialog(self, file_type):
        """Show dialog to enter custom filename"""
        content = BoxLayout(orientation='vertical', padding=10, spacing=10)
        
        content.add_widget(Label(
            text='Enter filename for Word document:',
            size_hint=(1, 0.2)
        ))
        
        # Default filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        default_name = f'photo_grid_{timestamp}'
        
        filename_input = TextInput(
            text=default_name,
            multiline=False,
            size_hint=(1, 0.2)
        )
        content.add_widget(filename_input)
        
        # Info label
        content.add_widget(Label(
            text='Will be saved to Documents folder\nas: [filename].docx',
            size_hint=(1, 0.2),
            color=(0.7, 0.7, 0.7, 1)
        ))
        
        button_layout = BoxLayout(size_hint=(1, 0.2), spacing=10)
        save_btn = Button(text='Save', background_color=(0.2, 0.8, 0.2, 1))
        cancel_btn = Button(text='Cancel')
        button_layout.add_widget(save_btn)
        button_layout.add_widget(cancel_btn)
        content.add_widget(button_layout)
        
        popup = Popup(title='Save as Word', content=content, size_hint=(0.9, 0.5))
        
        def on_save(instance):
            filename = filename_input.text.strip()
            if not filename:
                filename = default_name
            
            # Remove extension if user added it
            if filename.endswith('.docx'):
                filename = filename[:-5]
            
            popup.dismiss()
            self.save_word_file(filename)
        
        def on_cancel(instance):
            popup.dismiss()
        
        save_btn.bind(on_press=on_save)
        cancel_btn.bind(on_press=on_cancel)
        
        popup.open()
    
    def save_word_file(self, filename):
        """Save Word document with given filename"""
        # Save to Documents folder
        if platform == 'android':
            from android.storage import primary_external_storage_path
            docs_path = os.path.join(primary_external_storage_path(), 'Documents')
        else:
            docs_path = os.path.join(self.storage_path, 'Documents')
        
        os.makedirs(docs_path, exist_ok=True)
        
        full_path = os.path.join(docs_path, f'{filename}.docx')
        
        try:
            self.generate_word_document(full_path)
            self.show_message('Success', f'Word document saved!\n\n{filename}.docx\n\nLocation: Documents folder')
        except Exception as e:
            self.show_message('Error', f'Failed to save:\n{str(e)}')
    

    def generate_word_document(self, filename):
        """Generate Word document with photos - BULLETPROOF VERSION"""
        try:
            # Step 1: Determine layout from CURRENT UI state (grid_spinner.text)
            # This ensures we use what the user sees, not a stale variable
            current_layout = self.grid_spinner.text if hasattr(self, 'grid_spinner') else self.grid_layout_type
            
            print(f"DEBUG: Saving with layout: {current_layout}")
            print(f"DEBUG: grid_layout_type: {self.grid_layout_type}")
            print(f"DEBUG: photos_per_page: {self.photos_per_page}")
            
            if current_layout == '2x1':
                rows, cols = 2, 1
                photos_per_page = 2
                # Dimensions that FIT on one page (tested and verified)
                col_width = Inches(7.5)
                row_height = Inches(4.0)  # Fits on page
                max_img_width = Inches(7.0)
                max_img_height = Inches(3.4)  # Fits on page
            else:  # 2x2
                rows, cols = 2, 2
                photos_per_page = 4
                # Match desktop version - tighter spacing for 2x2
                col_width = Inches(3.55)
                row_height = Inches(4.0)  # Reduced from 4.5 for tighter spacing
                max_img_width = Inches(3.3)
                max_img_height = Inches(3.4)  # Reduced from 3.8
            
            print(f"DEBUG: Using photos_per_page: {photos_per_page}, rows: {rows}, cols: {cols}")
            
            # Step 2: Ensure ALL photos are distributed to pages
            # BUT preserve existing page_photos and page_captions (don't overwrite!)
            # This is critical - only fill in missing photos, don't reorganize existing ones
            for idx, photo_path in enumerate(self.photos):
                page_num = idx // photos_per_page
                grid_idx = idx % photos_per_page
                page_key = f'{page_num}_{grid_idx}'
                
                # Only add if not already assigned (preserve manual assignments and captions)
                if page_key not in self.page_photos:
                    self.page_photos[page_key] = photo_path
                    
                    # Only set default caption if not already set
                    if page_key not in self.page_captions:
                        self.page_captions[page_key] = f'{self.caption_prefix} {idx + 1}'
            
            # Step 3: Create document
            doc = Document()
            
            # Step 4: Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.8)
                section.bottom_margin = Inches(0.8)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                
                # Add header safely
                try:
                    if self.header_text and self.header_text.strip():
                        header = section.header
                        if header.paragraphs:
                            header_para = header.paragraphs[0]
                            header_run = header_para.add_run(self.header_text)
                            header_run.font.size = Pt(10)
                            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except:
                    pass  # Skip header if error
                
                # Add footer safely
                try:
                    footer = section.footer
                    if footer.paragraphs:
                        footer.paragraphs[0].text = ''
                    
                    if self.footer_line1 and self.footer_line1.strip():
                        footer_para1 = footer.add_paragraph()
                        footer_run1 = footer_para1.add_run(self.footer_line1)
                        footer_run1.font.size = Pt(9)
                        footer_para1.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    if self.footer_line2 and self.footer_line2.strip():
                        footer_para2 = footer.add_paragraph()
                        footer_run2 = footer_para2.add_run(self.footer_line2)
                        footer_run2.font.size = Pt(9)
                        footer_para2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except:
                    pass  # Skip footer if error
            
            # Step 5: Get all page keys safely
            if not self.page_photos:
                doc.save(filename)
                return
            
            page_keys = sorted(self.page_photos.keys(), 
                             key=lambda x: (int(x.split('_')[0]), int(x.split('_')[1])))
            
            if not page_keys:
                doc.save(filename)
                return
            
            max_page = max([int(k.split('_')[0]) for k in page_keys])
            
            # Step 6: Generate ALL pages (not just current page)
            for page_num in range(max_page + 1):
                # Collect photos for this page using the CORRECT photos_per_page
                page_photos_list = []
                for i in range(photos_per_page):
                    page_key = f'{page_num}_{i}'
                    if page_key in self.page_photos:
                        page_photos_list.append((i, self.page_photos[page_key]))
                
                if not page_photos_list:
                    continue
                
                # Add title with layout-specific spacing (match desktop)
                title = doc.add_paragraph()
                page_title_text = self.page_titles.get(page_num, 'Title')
                title_run = title.add_run(page_title_text)
                title_run.font.size = Pt(self.title_font_size)
                title_run.font.bold = self.title_bold
                title_run.font.underline = self.title_underline
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Spacing that FITS on one page (tested and verified)
                if current_layout == '2x1':
                    title.space_after = Pt(2)
                    title.space_before = Pt(0)
                else:
                    # Compact spacing for 2x2 layout
                    title.space_after = Pt(4)
                    title.space_before = Pt(0)
                
                # Create table
                table = doc.add_table(rows=rows, cols=cols)
                table.autofit = False
                table.allow_autofit = False  # Match desktop
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for col in table.columns:
                    col.width = col_width
                
                for row in table.rows:
                    row.height = row_height
                
                # Add photos to table
                for grid_idx, photo_path in page_photos_list:
                    try:
                        row_idx = grid_idx // cols
                        col_idx = grid_idx % cols
                        
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.vertical_alignment = 1
                        
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.space_before = Pt(0)
                        paragraph.space_after = Pt(0)
                        
                        # Add image
                        try:
                            img = Image.open(photo_path)
                            img_ratio = img.width / img.height
                            
                            if img_ratio > (max_img_width.inches / max_img_height.inches):
                                use_width = max_img_width
                                use_height = None
                            else:
                                use_width = None
                                use_height = max_img_height
                            
                            compressed_img = self.compress_image(photo_path)
                            run = paragraph.add_run()
                            if use_height:
                                run.add_picture(compressed_img, height=use_height)
                            else:
                                run.add_picture(compressed_img, width=use_width)
                        except:
                            paragraph.add_run(f'[Error: {os.path.basename(photo_path)}]')
                        
                        # Add caption with layout-specific spacing (match desktop)
                        caption_para = cell.add_paragraph()
                        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # EXACT copy from desktop version
                        if current_layout == '2x1':
                            caption_para.space_before = Pt(2)  # Tighter spacing
                            caption_para.space_after = Pt(0)
                        else:
                            # Tighter spacing for captions in 2x2 layout
                            caption_para.space_before = Pt(2)  # Reduced from 3
                            caption_para.space_after = Pt(0)
                        
                        page_key = f'{page_num}_{grid_idx}'
                        # Use saved caption or calculate sequential number
                        if page_key in self.page_captions:
                            caption_text = self.page_captions[page_key]
                        else:
                            # Calculate sequential photo number across all pages
                            photo_num = page_num * photos_per_page + grid_idx + 1
                            caption_text = f'{self.caption_prefix} {photo_num}'
                        
                        caption_run = caption_para.add_run(caption_text)
                        caption_run.font.size = Pt(self.caption_font_size)
                        caption_run.font.bold = self.caption_bold
                        caption_run.font.underline = self.caption_underline
                    except Exception as e:
                        print(f'Error adding photo {grid_idx}: {e}')
                        continue
                
                # Add page break if not last page
                # IMPORTANT: Don't add extra paragraphs or spacing after the table
                if page_num < max_page:
                    doc.add_page_break()
            
            # Step 7: Save document
            doc.save(filename)
            
        except Exception as e:
            print(f'Error generating document: {e}')
            raise
    
    def compress_image(self, image_path):
        """Compress image for document"""
        img = Image.open(image_path)
        
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        
        max_size = (1600, 1600)
        img.thumbnail(max_size, Image.Resampling.LANCZOS)
        
        output = io.BytesIO()
        img.save(output, format='JPEG', quality=92, optimize=True)
        output.seek(0)
        
        return output
    
    def show_message(self, title, message):
        """Show popup message"""
        content = BoxLayout(orientation='vertical', padding=10, spacing=10)
        content.add_widget(Label(text=message))
        
        ok_btn = Button(text='OK', size_hint=(1, 0.3))
        content.add_widget(ok_btn)
        
        popup = Popup(title=title, content=content, size_hint=(0.8, 0.4))
        ok_btn.bind(on_press=popup.dismiss)
        popup.open()
    
    def save_settings(self):
        """Save settings to JSON"""
        try:
            settings = {
                'caption_prefix': self.caption_prefix,
                'header_text': self.header_text,
                'footer_line1': self.footer_line1,
                'footer_line2': self.footer_line2,
                'grid_layout': self.grid_layout_type,
                'last_photo_path': self.last_photo_path,
                'title_font_size': self.title_font_size,
                'title_bold': self.title_bold,
                'title_underline': self.title_underline,
                'caption_font_size': self.caption_font_size,
                'caption_bold': self.caption_bold,
                'caption_underline': self.caption_underline
            }
            with open(self.settings_file, 'w') as f:
                json.dump(settings, f, indent=2)
        except Exception as e:
            print(f'Failed to save settings: {e}')
    
    def load_settings(self):
        """Load settings from JSON"""
        try:
            if os.path.exists(self.settings_file):
                with open(self.settings_file, 'r') as f:
                    settings = json.load(f)
                
                self.caption_prefix = settings.get('caption_prefix', 'Photo')
                self.header_text = settings.get('header_text', '')
                self.footer_line1 = settings.get('footer_line1', 'File Ref. :')
                self.footer_line2 = settings.get('footer_line2', f'Photo taken on {datetime.now().strftime("%d.%m.%Y")}')
                
                # Load grid layout and sync photos_per_page
                # Always default to 2x2 for new sessions
                self.grid_layout_type = settings.get('grid_layout', '2x2')
                # Override to 2x2 if it was 2x1 (force 2x2 as default)
                if self.grid_layout_type == '2x1':
                    self.grid_layout_type = '2x2'  # Force 2x2 default
                
                if self.grid_layout_type == '2x1':
                    self.photos_per_page = 2
                else:
                    self.photos_per_page = 4
                
                self.last_photo_path = settings.get('last_photo_path', self.storage_path)
                self.title_font_size = settings.get('title_font_size', 14)
                self.title_bold = settings.get('title_bold', True)
                self.title_underline = settings.get('title_underline', True)
                self.caption_font_size = settings.get('caption_font_size', 10)
                self.caption_bold = settings.get('caption_bold', False)
                self.caption_underline = settings.get('caption_underline', False)
                
                print(f"DEBUG: Loaded settings - grid_layout_type: {self.grid_layout_type}, photos_per_page: {self.photos_per_page}")
        except Exception as e:
            print(f'Failed to load settings: {e}')


if __name__ == '__main__':
    PhotoGridApp().run()
